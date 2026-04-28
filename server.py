"""
IDIEM — Servidor local para generación de propuestas Word
=========================================================
Cómo usar:
1. Instala dependencias:  pip install flask flask-cors python-docx anthropic python-dotenv
2. Crea archivo .env con: ANTHROPIC_API_KEY=tu_key_aqui
3. Corre el servidor:     python server.py
4. Abre el HTML en Chrome y genera propuestas normalmente

El servidor corre en http://localhost:5050
No cierres esta ventana mientras usas el generador.

Cuando quieras deployar en DigitalOcean App Platform:
- Sube el repo (sin .env)
- Configura ANTHROPIC_API_KEY como variable de entorno en el panel de DO
- Cambia SERVER_URL en el HTML a tu URL de DO
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv
from openai import OpenAI
import io
import os
import json
import re
from datetime import datetime

load_dotenv()

app = Flask(__name__)
CORS(app)

# ============================================================
# CONFIGURACION
# ============================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Cliente usando Serverless Inference de DigitalOcean
# Sin limite de tokens, sin necesitar key de Anthropic
# Agrega esto a tu .env:
#   DO_INFERENCE_KEY=sk-do-XXXX...  (la Model Access Key que ya usaste)
client = OpenAI(
    base_url="https://inference.do-ai.run/v1",
    api_key=os.getenv("DO_INFERENCE_KEY", "missing-key")
)

# Modelo Llama 3.3 hosteado por DO
# Alternativas disponibles sin key externa:
#   "deepseek-r1-distill-llama-70b"
#   "minimax-m2.5"
MODEL = "llama3.3-70b-instruct"

TEMPLATES = {
    'diagnostico': os.path.join(BASE_DIR, 'plantillas', 'template_diagnostico_IDIEM.docx'),
    'calculo':     os.path.join(BASE_DIR, 'plantillas', 'template_diagnostico_IDIEM.docx'),  # usa mismo template por ahora
    'incendios':   os.path.join(BASE_DIR, 'plantillas', 'template_diagnostico_IDIEM.docx'),  # usa mismo template por ahora
    # 'revision':  os.path.join(BASE_DIR, 'plantillas', 'template_revision_IDIEM.docx'),
    # 'monitoreo': os.path.join(BASE_DIR, 'plantillas', 'template_monitoreo_IDIEM.docx'),
}

PROMPTS = {
    'diagnostico': os.path.join(BASE_DIR, 'prompts', 'diagnostico.txt'),
    'calculo':     os.path.join(BASE_DIR, 'prompts', 'calculo.txt'),
    'incendios':   os.path.join(BASE_DIR, 'prompts', 'incendios.txt'),
    # 'revision':  os.path.join(BASE_DIR, 'prompts', 'revision.txt'),
    # 'monitoreo': os.path.join(BASE_DIR, 'prompts', 'monitoreo.txt'),
}

MARKERS = {
    'alcance':             '▶ PEGAR AQUÍ: SECCIÓN 1 — ALCANCE (texto generado por el generador HTML)',
    'propuesta_tecnica':   '▶ PEGAR AQUÍ: SECCIÓN 2 — PROPUESTA TÉCNICA (texto generado por el generador HTML)',
    'plazos':              '▶ PEGAR AQUÍ: SECCIÓN 3 — PLAZOS (texto generado por el generador HTML)',
    'propuesta_economica': '▶ PEGAR AQUÍ: SECCIÓN 4 — PROPUESTA ECONÓMICA (texto generado por el generador HTML)',
    'recursos':            '▶ PEGAR AQUÍ: SECCIÓN 5 — RECURSOS DEL MANDANTE (texto generado por el generador HTML)',
    'exclusiones':         '▶ PEGAR AQUÍ: SECCIÓN 6 — EXCLUSIONES (texto generado por el generador HTML)',
}

# ============================================================
# LIMPIEZA DE DATOS DEL FORMULARIO
# ============================================================

def limpiar_nombre(nombre):
    """
    Convierte nombres a formato título correcto.
    'GUILLERMO SIERRA' -> 'Guillermo Sierra'
    'geneSISs diaaz' -> 'Génesis Díaz'
    """
    if not nombre:
        return nombre
    # Quitar puntuación extra al final
    nombre = re.sub(r'[.,;:\s]+$', '', nombre.strip())
    # Convertir a título
    return nombre.title()


def limpiar_datos(datos):
    """
    Limpia y normaliza los datos del formulario antes de
    enviarlos al modelo. Esto hace el sistema robusto a
    errores de tipeo, mayúsculas incorrectas, etc.
    """
    limpio = datos.copy()

    # Normalizar nombres de personas
    for campo in ['elaborado', 'revisado', 'contacto']:
        if limpio.get(campo):
            limpio[campo] = limpiar_nombre(limpio[campo])

    # Limpiar espacios extra en todos los campos de texto
    for campo, valor in limpio.items():
        if isinstance(valor, str):
            limpio[campo] = ' '.join(valor.split())

    # Valores por defecto para campos vacíos críticos
    defaults = {
        'motivacion':   'Evaluación del estado estructural actual del inmueble.',
        'obs_tecnicas': 'Sin observaciones adicionales.',
        'plazo_terreno':  'A definir',
        'plazo_gabinete': 'A definir',
        'plazo_total':    'A definir',
    }
    for campo, default in defaults.items():
        if not limpio.get(campo) or limpio[campo].strip() == '':
            limpio[campo] = default

    return limpio


# ============================================================
# CARGAR Y CONSTRUIR PROMPT
# ============================================================

def cargar_prompt(tipo, datos_formulario):
    """
    Lee el archivo de prompt, limpia los datos del formulario
    y reemplaza los placeholders con los valores reales.
    """
    prompt_path = PROMPTS.get(tipo)
    if not prompt_path or not os.path.exists(prompt_path):
        raise FileNotFoundError(f'Prompt no encontrado: {prompt_path}')

    with open(prompt_path, 'r', encoding='utf-8') as f:
        prompt = f.read()

    # Limpiar datos antes de insertarlos en el prompt
    datos = limpiar_datos(datos_formulario)

    # Forma de pago legible
    forma_map = {
        '50_inicio':         '50% al momento de aceptar el servicio, 50% al entregar el informe',
        '50_terreno':        '50% al término de los trabajos en terreno, 50% al entregar el informe',
        'estados_mensuales': 'Estados de pago mensuales',
        'otro':              'Según detalle'
    }
    forma_pago = forma_map.get(datos.get('forma_pago', '50_inicio'), '')
    if datos.get('detalle_pago'):
        forma_pago += '. ' + datos['detalle_pago']

    actividades = datos.get('actividades', [])
    actividades_str = '\n'.join(f'- {a}' for a in actividades) if actividades else '- No especificadas'

    reemplazos = {
        # Campos compartidos
        '{{cliente}}':         datos.get('cliente', 'No indicado'),
        '{{contacto}}':        datos.get('contacto', 'No indicado'),
        '{{direccion}}':       datos.get('direccion', 'No indicada'),
        '{{superficie}}':      datos.get('superficie', 'No indicada'),
        '{{motivacion}}':      datos.get('motivacion', 'No indicada'),
        '{{obs_tecnicas}}':    datos.get('obs_tecnicas', 'Ninguna'),
        '{{actividades}}':     actividades_str,
        '{{plazo_terreno}}':   datos.get('plazo_terreno', 'No indicado'),
        '{{plazo_gabinete}}':  datos.get('plazo_gabinete', 'No indicado'),
        '{{plazo_total}}':     datos.get('plazo_total', 'No indicado'),
        '{{precio_uf}}':       datos.get('precio_uf', 'A definir'),
        '{{forma_pago}}':      forma_pago,
        '{{exclusiones}}':     datos.get('exclusiones', 'Ninguna adicional'),
        '{{codigo}}':          datos.get('codigo', 'PR.DEP.2025.XXXX'),
        '{{revision}}':        datos.get('revision', 'Rev. N°1'),
        '{{fecha}}':           datos.get('fecha', datetime.now().strftime('%d-%m-%Y')),
        '{{elaborado}}':       datos.get('elaborado', 'Pendiente'),
        '{{revisado}}':        datos.get('revisado', 'Guillermo Sierra R.'),
        # Campos diagnostico
        '{{edificio}}':        datos.get('edificio', datos.get('proyecto', 'No indicado')),
        '{{anio}}':            datos.get('anio', 'No indicado'),
        '{{pisos}}':           datos.get('pisos', 'No indicado'),
        '{{materialidad}}':    datos.get('materialidad', datos.get('tipo_estructura', 'No especificada')),
        '{{condicion}}':       datos.get('condicion', 'No especificada'),
        # Campos calculo
        '{{tipo_calculo}}':    datos.get('tipo_calculo', 'proyecto nuevo'),
        '{{proyecto}}':        datos.get('proyecto', datos.get('edificio', 'No indicado')),
        '{{tipo_estructura}}': datos.get('tipo_estructura', datos.get('materialidad', 'No especificada')),
        '{{uso}}':             datos.get('uso', 'No indicado'),
        '{{informe_previo}}':  datos.get('informe_previo', 'No aplica'),
        '{{normativa}}':       datos.get('normativa', 'Normativa chilena vigente'),
        # Campos incendios
        '{{fecha_incendio}}':   datos.get('fecha_incendio', 'No indicada'),
        '{{descripcion_siniestro}}': datos.get('descripcion_siniestro', 'No indicada'),
        '{{plazo_preliminar}}': datos.get('plazo_preliminar', 'No aplica'),
    }

    for placeholder, valor in reemplazos.items():
        prompt = prompt.replace(placeholder, str(valor))

    return prompt


# ============================================================
# LLAMADA A CLAUDE
# ============================================================

def llamar_claude(prompt):
    """
    Llama al agente de DigitalOcean via OpenAI-compatible API.
    El agente tiene Claude configurado internamente en DO.
    """
    response = client.chat.completions.create(
        model=MODEL,
        max_tokens=4096,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )
    return response.choices[0].message.content


# ============================================================
# PARSEAR JSON DE CLAUDE
# ============================================================

def parsear_json(texto_crudo):
    """
    Extrae el JSON de la respuesta del modelo de forma robusta.
    Maneja: markdown, JSON sin llaves externas, escape doble de \n.
    """
    texto = re.sub(r'```json\s*', '', texto_crudo)
    texto = re.sub(r'```\s*', '', texto)
    texto = texto.strip()

    def sanitizar_json(t):
        """Escapa saltos de línea literales dentro de strings JSON."""
        resultado = []
        dentro_string = False
        escape = False
        for ch in t:
            if escape:
                resultado.append(ch)
                escape = False
            elif ch == '\\':
                resultado.append(ch)
                escape = True
            elif ch == '"':
                dentro_string = not dentro_string
                resultado.append(ch)
            elif dentro_string and ch == '\n':
                resultado.append('\\n')
            elif dentro_string and ch == '\r':
                resultado.append('\\r')
            elif dentro_string and ch == '\t':
                resultado.append('\\t')
            else:
                resultado.append(ch)
        return ''.join(resultado)

    # Intento 1: parseo directo
    try:
        data = json.loads(texto)
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass

    # Intento 2: sanitizar saltos de línea literales y parsear
    try:
        data = json.loads(sanitizar_json(texto))
        if isinstance(data, dict):
            return data
    except json.JSONDecodeError:
        pass

    # Intento 3: buscar el objeto JSON entre { } y sanitizar
    inicio = texto.find('{')
    fin    = texto.rfind('}')
    if inicio != -1 and fin != -1 and fin > inicio:
        try:
            data = json.loads(sanitizar_json(texto[inicio:fin+1]))
            if isinstance(data, dict):
                return data
        except json.JSONDecodeError:
            pass

    # Intento 4: extraer campo por campo con regex como fallback
    campos = ['alcance', 'propuesta_tecnica', 'plazos', 'propuesta_economica', 'recursos', 'exclusiones']
    resultado = {}
    for campo in campos:
        patron = rf'"{campo}"\s*:\s*"((?:[^"\\]|\\.)*)"' 
        m = re.search(patron, texto, re.DOTALL)
        if m:
            resultado[campo] = m.group(1).replace('\\n', '\n').replace('\\"', '"'). replace('\\\\', '\\')
    if not resultado:
        # Intento 5: regex con DOTALL para capturar saltos de línea literales
        for campo in campos:
            patron = rf'"{campo}"\s*:\s*"(.*?)(?<!\\)"'
            m = re.search(patron, texto, re.DOTALL)
            if m:
                resultado[campo] = m.group(1).strip()
    if resultado:
        print(f'  → JSON extraído por regex ({len(resultado)} campos)')
        return resultado

    print(f'⚠ No se pudo parsear JSON. Primeros 500 chars:\n{texto[:500]}')
    return {
        'alcance':             texto,
        'propuesta_tecnica':   '',
        'plazos':              '',
        'propuesta_economica': '',
        'recursos':            '',
        'exclusiones':         ''
    }


# ============================================================
# RELLENAR WORD
# ============================================================

def detectar_estilo_linea(linea, doc):
    """
    Detecta el estilo Word a aplicar según el contenido de la línea.
    Usa los estilos definidos en el template de IDIEM.
    """
    import re
    linea = linea.strip()

    # Subtítulo numerado: 2.1, 2.2, 2.5.1, ETAPA 1:, etc.
    if re.match(r'^\d+\.\d+', linea) or re.match(r'^ETAPA \d+', linea, re.IGNORECASE):
        estilo = 'Subtitulo 2'
    # Lista con guión o bala
    elif linea.startswith('- ') or linea.startswith('• '):
        estilo = 'cuerpo 1 espacio'
    # Tabla (contiene |)
    elif '|' in linea:
        estilo = 'texto tablas'
    # Línea vacía
    elif not linea:
        estilo = 'Normal'
    # Texto normal
    else:
        estilo = 'Cuerpo'

    # Verificar que el estilo existe en el documento
    estilos_disponibles = [s.name for s in doc.styles]
    if estilo not in estilos_disponibles:
        estilo = 'Normal'

    return estilo


def insertar_tabla_simple(doc, para_ref, lineas_tabla):
    """
    Convierte líneas con | en una tabla Word real.
    Formato: 'Columna A | Columna B'
    """
    from docx.oxml.ns import qn
    import re

    filas_datos = []
    for linea in lineas_tabla:
        if '|' in linea:
            celdas = [c.strip() for c in linea.split('|') if c.strip()]
            if celdas:
                filas_datos.append(celdas)

    if not filas_datos:
        return

    max_cols = max(len(f) for f in filas_datos)
    tabla = doc.add_table(rows=len(filas_datos), cols=max_cols)
    tabla.style = 'Table Grid'

    for ri, fila in enumerate(filas_datos):
        for ci, texto in enumerate(fila):
            if ci < max_cols:
                cell = tabla.cell(ri, ci)
                cell.text = texto
                # Aplicar estilo de tabla
                for p in cell.paragraphs:
                    try:
                        p.style = doc.styles['texto tablas']
                    except Exception:
                        pass
                # Primera fila en negrita
                if ri == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True

    # Mover la tabla al lugar del párrafo de referencia
    para_ref._element.addprevious(tabla._tbl)


def reemplazar_marcador(doc, marcador, contenido):
    """
    Reemplaza un marcador rojo en el documento con el contenido generado.
    Aplica estilos automáticamente según el tipo de línea:
    - Subtítulos 2.1, 2.2 → estilo 'Subtitulo 2'
    - Listas con guión → estilo 'cuerpo 1 espacio'
    - Texto normal → estilo 'Cuerpo'
    - Líneas con | → tabla Word real
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy, re

    for para in doc.paragraphs:
        if marcador not in para.text:
            continue

        para_element = para._element
        parent = para_element.getparent()
        idx = list(parent).index(para_element)

        # Limpiar el párrafo marcador
        for run in para.runs:
            run.text = ''

        lineas = contenido.split('\n')

        # Agrupar líneas de tabla para procesarlas juntas
        i = 0
        parrafos_insertados = 0

        while i < len(lineas):
            linea = lineas[i]
            linea_stripped = linea.strip()

            # Detectar bloque de tabla (líneas consecutivas con |)
            if '|' in linea_stripped:
                bloque_tabla = []
                while i < len(lineas) and '|' in lineas[i]:
                    bloque_tabla.append(lineas[i])
                    i += 1
                # Insertar tabla antes del marcador
                insertar_tabla_simple(doc, para, bloque_tabla)
                continue

            # Detectar estilo según contenido
            estilo = detectar_estilo_linea(linea_stripped, doc)

            if i == 0:
                # Primera línea: usar el párrafo marcador
                try:
                    para.style = doc.styles[estilo]
                except Exception:
                    pass
                if linea_stripped:
                    run = para.add_run(linea_stripped)
                    run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                # Líneas siguientes: nuevos párrafos con estilo
                new_para = OxmlElement('w:p')

                # Aplicar estilo al nuevo párrafo
                pPr = OxmlElement('w:pPr')
                pStyle = OxmlElement('w:pStyle')
                # Buscar ID del estilo
                estilo_id = estilo
                for s in doc.styles:
                    if s.name == estilo:
                        estilo_id = s.style_id
                        break
                pStyle.set(qn('w:val'), estilo_id)
                pPr.append(pStyle)
                new_para.append(pPr)

                if linea_stripped:
                    new_r = OxmlElement('w:r')
                    new_rPr = OxmlElement('w:rPr')
                    new_color = OxmlElement('w:color')
                    new_color.set(qn('w:val'), '000000')
                    new_rPr.append(new_color)
                    new_r.append(new_rPr)
                    new_t = OxmlElement('w:t')
                    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    new_t.text = linea_stripped
                    new_r.append(new_t)
                    new_para.append(new_r)

                parent.insert(idx + parrafos_insertados + 1, new_para)
                parrafos_insertados += 1

            i += 1

        return True

    return False


def extraer_numero_revision(revision_str):
    """
    De 'Rev. N°2', 'Revisión A', 'B', '3' → extrae solo '2', 'A', 'B', '3'.
    """
    if not revision_str:
        return '1'
    # Buscar letra o número al final
    import re
    m = re.search(r'([A-Za-z0-9]+)\s*$', revision_str.strip())
    return m.group(1) if m else revision_str.strip()


def insertar_campo_numpages(para):
    """
    Reemplaza el contenido del párrafo con un campo NUMPAGES de Word.
    Word lo actualiza automáticamente al abrir o imprimir el documento.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Preservar propiedades del párrafo (rPr) del primer run si existe
    rPr_orig = None
    if para.runs:
        rPr_orig = para.runs[0]._element.find(qn('w:rPr'))

    # Borrar todos los runs existentes
    for r in para._element.findall(qn('w:r')):
        para._element.remove(r)

    def make_run(child):
        r = OxmlElement('w:r')
        if rPr_orig is not None:
            import copy
            r.append(copy.deepcopy(rPr_orig))
        r.append(child)
        return r

    # Campo Word: fldChar begin → instrText → fldChar end
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' NUMPAGES '

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    para._element.append(make_run(fld_begin))
    para._element.append(make_run(instr))
    para._element.append(make_run(fld_end))


def reemplazar_marcador_portada(cell_o_parrafos, marcador, valor):
    """
    Reemplaza un marcador en párrafos de una celda.
    Maneja runs divididos consolidando todo el texto en un run.
    Si el valor contiene \n, inserta párrafos adicionales.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    parrafos = cell_o_parrafos.paragraphs if hasattr(cell_o_parrafos, 'paragraphs') else cell_o_parrafos

    for para in parrafos:
        if marcador not in para.text:
            continue

        texto_original = para.text
        valor_str = str(valor)
        texto_nuevo = texto_original.replace(marcador, valor_str)

        # Guardar propiedades del primer run (fuente, tamaño, negrita, etc.)
        primer_run_xml = None
        if para.runs:
            primer_run_xml = copy.deepcopy(para.runs[0]._element)

        # Borrar todos los runs existentes
        for r in para._element.findall(qn('w:r')):
            para._element.remove(r)

        lineas = texto_nuevo.split('\n')

        # Primera línea en el párrafo actual
        r_nuevo = OxmlElement('w:r')
        if primer_run_xml is not None:
            # Copiar propiedades del run original (rPr)
            rPr_orig = primer_run_xml.find(qn('w:rPr'))
            if rPr_orig is not None:
                r_nuevo.append(copy.deepcopy(rPr_orig))
        t_nuevo = OxmlElement('w:t')
        t_nuevo.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t_nuevo.text = lineas[0]
        r_nuevo.append(t_nuevo)
        para._element.append(r_nuevo)

        # Líneas adicionales: nuevos párrafos con misma estructura
        if len(lineas) > 1:
            parent = para._element.getparent()
            idx = list(parent).index(para._element)
            for i, linea in enumerate(lineas[1:], 1):
                nuevo_p = copy.deepcopy(para._element)
                # Limpiar runs del nuevo párrafo
                for r in nuevo_p.findall(qn('w:r')):
                    nuevo_p.remove(r)
                # Agregar run con el texto
                r_elem = OxmlElement('w:r')
                if primer_run_xml is not None:
                    rPr_orig = primer_run_xml.find(qn('w:rPr'))
                    if rPr_orig is not None:
                        r_elem.append(copy.deepcopy(rPr_orig))
                t_elem = OxmlElement('w:t')
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t_elem.text = linea
                r_elem.append(t_elem)
                nuevo_p.append(r_elem)
                parent.insert(idx + i, nuevo_p)
        break


def actualizar_portada(doc, datos):
    """
    Actualiza todos los marcadores de la portada con datos del formulario.
    Marcadores soportados: {{TITULO}}, {{PR_CODE}}, {{REVISION}},
    {{N_PAGINAS}}, {{FECHA}}, {{ELABORADO}}, {{REVISADO}}, {{CLIENTE}}
    """
    datos_limpios = limpiar_datos(datos)

    # Construir título del proyecto
    edificio = datos_limpios.get('edificio') or datos_limpios.get('proyecto', '')
    cliente  = datos_limpios.get('cliente', '')
    titulo   = f"Diagnóstico Estructural {edificio}" if edificio else "Propuesta Técnica y Económica"

    # Revisión — solo el número o letra
    revision_completa = datos_limpios.get('revision', 'Rev. N°1')
    revision_corta    = extraer_numero_revision(revision_completa)

    # N_PAGINAS se maneja con campo dinámico Word — no se incluye en marcadores normales

    # Elaborado y revisado — pueden ser múltiples nombres separados por coma o /
    elaborado_raw = datos_limpios.get('elaborado', '')
    revisado_raw  = datos_limpios.get('revisado', 'Guillermo Sierra R.')

    # Normalizar separadores → \n para inserción multilínea
    import re
    elaborado = re.sub(r'[,/]+', '\n', elaborado_raw).strip()
    revisado  = re.sub(r'[,/]+', '\n', revisado_raw).strip()

    marcadores = {
        '{{TITULO}}':    titulo,
        '{{PR_CODE}}':   datos_limpios.get('codigo', 'PR.DEP.2025.XXXX'),
        '{{REVISION}}':  revision_corta,
        '{{FECHA}}':     datos_limpios.get('fecha', datetime.now().strftime('%d-%m-%Y')),
        '{{ELABORADO}}': elaborado,
        '{{REVISADO}}':  revisado,
        '{{CLIENTE}}':   cliente,
    }

    # Reemplazar en párrafos sueltos
    for para in doc.paragraphs:
        # Campo dinámico NUMPAGES para {{N_PAGINAS}}
        if '{{N_PAGINAS}}' in para.text:
            insertar_campo_numpages(para)
            continue
        for marcador, valor in marcadores.items():
            if marcador in para.text:
                reemplazar_marcador_portada(para, marcador, valor)

    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Campo dinámico NUMPAGES
                if '{{N_PAGINAS}}' in cell._element.xml:
                    for para in cell.paragraphs:
                        if '{{N_PAGINAS}}' in para.text:
                            insertar_campo_numpages(para)
                    continue
                for marcador, valor in marcadores.items():
                    if marcador in cell._element.xml:
                        reemplazar_marcador_portada(cell, marcador, valor)


# ============================================================
# RUTAS
# ============================================================

@app.route('/health', methods=['GET'])
def health():
    key_ok = bool(os.getenv("DO_INFERENCE_KEY"))
    return jsonify({
        'status':  'ok',
        'endpoint': 'https://inference.do-ai.run/v1',
        'api_key': 'configurada' if key_ok else 'FALTA — agrega DO_INFERENCE_KEY al .env',
        'modelo':  MODEL
    })


@app.route('/generar', methods=['POST'])
def generar():
    """
    Endpoint principal — recibe datos del formulario,
    llama a Claude, parsea el JSON, rellena el Word y lo devuelve.
    Todo en un solo paso desde el punto de vista del HTML.
    """
    try:
        data       = request.get_json()
        tipo       = data.get('tipo', 'diagnostico')
        datos_form = data.get('formulario', {})

        # Verificar template
        template_path = TEMPLATES.get(tipo)
        if not template_path or not os.path.exists(template_path):
            return jsonify({'error': f'Template no encontrado: {template_path}'}), 404

        # Construir prompt con datos limpios
        nombre_proyecto = datos_form.get('edificio') or datos_form.get('proyecto') or '?'
        print(f'\n[{datetime.now().strftime("%H:%M:%S")}] Generando propuesta: {nombre_proyecto}')
        prompt = cargar_prompt(tipo, datos_form)

        # Llamar a Claude
        print(f'  → Llamando a Claude ({MODEL})...')
        texto_crudo = llamar_claude(prompt)
        print(f'  → Respuesta recibida ({len(texto_crudo)} chars)')

        # Parsear JSON
        secciones = parsear_json(texto_crudo)
        for k, v in secciones.items():
            print(f'  → {k}: {len(v)} chars')

        # Rellenar Word
        doc = Document(template_path)
        actualizar_portada(doc, datos_form)

        for key, marcador in MARKERS.items():
            contenido = secciones.get(key, '')
            if contenido:
                ok = reemplazar_marcador(doc, marcador, contenido)
                if not ok:
                    print(f'  ⚠ Marcador no encontrado: {key}')
            else:
                print(f'  ⚠ Sección vacía: {key}')

        # Generar nombre y devolver archivo
        edificio = (datos_form.get('edificio') or datos_form.get('proyecto') or 'propuesta').replace(' ', '_')
        codigo   = datos_form.get('codigo', 'PR_DEP_2025_XXXX').replace('.', '_')
        nombre   = f'{codigo}_{edificio}.docx'

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        print(f'  ✓ Word generado: {nombre}')

        return send_file(
            buffer,
            as_attachment=True,
            download_name=nombre,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as auth_err:
        if '401' in str(auth_err) or 'authentication' in str(auth_err).lower():
            return jsonify({'error': 'Access key invalida. Verifica DO_INFERENCE_KEY en el .env'}), 401
        if '429' in str(auth_err):
            return jsonify({'error': 'Limite de uso alcanzado. Intenta en unos minutos.'}), 429
    except Exception as e:
        import traceback
        print(traceback.format_exc())
        return jsonify({'error': str(e), 'detalle': traceback.format_exc()}), 500


# ============================================================
# INICIO
# ============================================================
if __name__ == '__main__':
    key = os.getenv("DO_INFERENCE_KEY")
    print('\n' + '='*55)
    print('  IDIEM — Servidor de propuestas (DO Serverless)')
    print('='*55)
    print(f'  URL:      http://localhost:5050')
    print(f'  Endpoint: https://inference.do-ai.run/v1')
    print(f'  Modelo:   {MODEL}')
    print(f'  Key:      {"OK" if key else "FALTA — agrega DO_INFERENCE_KEY al .env"}')
    print('\n  Templates:')
    for key, path in TEMPLATES.items():
        print(f'    {"✓" if os.path.exists(path) else "✗ NO ENCONTRADO"}  {key}')
    print('\n  Prompts:')
    for key, path in PROMPTS.items():
        print(f'    {"✓" if os.path.exists(path) else "✗ NO ENCONTRADO"}  {key}')
    print('='*55)
    print('  Deja esta ventana abierta mientras usas el generador.')
    print('  Ctrl + C para detener.')
    print('='*55 + '\n')
    app.run(host='localhost', port=5050, debug=False)